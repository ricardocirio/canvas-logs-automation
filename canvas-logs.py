#!/usr/bin/env python3
"""
Export Canvas activity logs and submissions from PostgreSQL.

Usage:
  python canvas-logs.py --username <unique_id> --start "2025-08-26 00:00:00" --end "2025-10-31 00:00:00" --output-dir results

This will create the specified folder and generate:
  - activity.xlsx (web logs)
  - submissions.xlsx (assignment submissions with forensics)

Environment (PostgreSQL):
  POSTGRES_DSN         (optional, full DSN string)
  or the standard libpq vars:
  PGHOST, PGPORT, PGDATABASE, PGUSER, PGPASSWORD, PGSSLMODE
"""
from __future__ import annotations

import argparse
import os
import sys
import time
from datetime import datetime
from typing import Dict, Optional
from zoneinfo import ZoneInfo

try:
    import psycopg
except Exception as e:  # pragma: no cover
    print("Missing dependency: psycopg. Install with: pip install -r requirements.txt", file=sys.stderr)
    raise

try:
    import pandas as pd
except Exception as e:  # pragma: no cover
    print("Missing dependency: pandas. Install with: pip install -r requirements.txt", file=sys.stderr)
    raise

try:
    import requests
except Exception as e:  # pragma: no cover
    print("Missing dependency: requests. Install with: pip install -r requirements.txt", file=sys.stderr)
    raise

# Optional Word document reporting dependency
try:
    from docx import Document
    _DOCX_AVAILABLE = True
except Exception:
    _DOCX_AVAILABLE = False


REPO_DIR = os.path.dirname(os.path.abspath(__file__))

QUERY_TYPES = {
    "activity": os.path.join(REPO_DIR, "sql", "activity.postgres.sql"),
    "submissions": os.path.join(REPO_DIR, "sql", "submissions.postgres.sql"),
}

# Cache for IP geolocation lookups
IP_CACHE: Dict[str, Dict[str, Optional[str]]] = {}


def get_ip_location(ip: str) -> Dict[str, Optional[str]]:
    """Get location info for an IP address using ipinfo.io with fallback to ip-api.com.

    Returns a dict with keys: country, region, city, isp
    """
    if not ip:
        return {"country": None, "region": None, "city": None, "isp": None}
    
    # Check cache first
    if ip in IP_CACHE:
        return IP_CACHE[ip]
    
    location = {"country": None, "region": None, "city": None, "isp": None}
    
    # Try ipinfo.io first (rate limited, free tier: ~50k requests/month)
    try:
        response = requests.get(f"https://ipinfo.io/{ip}/json", timeout=5)
        if response.status_code == 200:
            data = response.json()
            location["country"] = data.get("country")
            location["region"] = data.get("region")
            location["city"] = data.get("city")
            # ipinfo returns ISP/organization under 'org' (e.g., 'AS1234 Some ISP')
            location["isp"] = data.get("org")
            IP_CACHE[ip] = location
            time.sleep(0.1)  # Rate limiting
            return location
    except Exception:
        pass
    
    # Fallback to ipwho.is (no API key, fair-use limits)
    try:
        response = requests.get(f"https://ipwho.is/{ip}", timeout=5)
        if response.status_code == 200:
            data = response.json()
            if data.get("success") is True:
                location["country"] = data.get("country_code")
                location["region"] = data.get("region")
                location["city"] = data.get("city")
                conn = data.get("connection") or {}
                location["isp"] = conn.get("isp") or conn.get("org")
                IP_CACHE[ip] = location
                time.sleep(0.1)  # Rate limiting
                return location
    except Exception:
        pass
    
    # Cache even if lookup failed to avoid repeated attempts
    IP_CACHE[ip] = location
    return location


def parse_dt(value: str) -> datetime:
    """Parse a timestamp string in EST. Accepts 'YYYY-MM-DD HH:MM:SS' or ISO 'YYYY-MM-DDTHH:MM:SS'."""
    value = value.strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(value, fmt)
            if fmt == "%Y-%m-%d":
                dt = datetime(dt.year, dt.month, dt.day, 0, 0, 0)
            # Attach EST timezone
            return dt.replace(tzinfo=ZoneInfo("America/New_York"))
        except ValueError:
            continue
    raise argparse.ArgumentTypeError(
        f"Invalid datetime: '{value}'. Use 'YYYY-MM-DD HH:MM:SS' or ISO 'YYYY-MM-DDTHH:MM:SS'."
    )


def read_sql(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def connect_postgres():
    dsn = os.environ.get("POSTGRES_DSN")
    if dsn:
        # Add client_encoding to handle older PostgreSQL servers
        if "client_encoding" not in dsn.lower():
            dsn += " client_encoding=utf8"
        return psycopg.connect(dsn)

    host = os.environ.get("PGHOST")
    dbname = os.environ.get("PGDATABASE")
    user = os.environ.get("PGUSER")
    password = os.environ.get("PGPASSWORD")
    port = os.environ.get("PGPORT", "5432")
    sslmode = os.environ.get("PGSSLMODE")

    missing = [k for k, v in (("PGHOST", host), ("PGDATABASE", dbname), ("PGUSER", user)) if not v]
    if missing:
        raise RuntimeError(f"Missing required environment variables: {', '.join(missing)} (or provide POSTGRES_DSN)")

    kwargs = dict(host=host, dbname=dbname, user=user, port=port, client_encoding="utf8")
    if password:
        kwargs["password"] = password
    if sslmode:
        kwargs["sslmode"] = sslmode
    return psycopg.connect(**kwargs)


# (Removed Markdown summary generation)

def write_submissions_summary_docx(
    df: pd.DataFrame,
    username: str,
    start_ts: datetime,
    end_ts: datetime,
    output_path: str,
) -> None:
    """Generate a Word (.docx) summary with the requested bullet format.

    Structure:
      <username> logs summary report
      Course Name: <course>
      Assignments Submitted:
      - <assignment>
        - Submitted: mm/dd/yyyy at hh:mm AM/PM
        - IP Address: <ip>
        - IP Address Location: city, region, country
    """
    if not _DOCX_AVAILABLE:
        print("  → Skipping Word summary (python-docx not installed)")
        return

    # Determine output DOCX path from the submissions Excel output path
    out_dir = os.path.dirname(output_path)
    dir_basename = os.path.basename(os.path.normpath(out_dir))
    docx_path = os.path.join(out_dir, f"{dir_basename}-summary.docx")

    # Columns expected
    assign_col = 'assignment' if 'assignment' in df.columns else None
    time_col = 'timestamp_est' if 'timestamp_est' in df.columns else (
        'submitted_at' if 'submitted_at' in df.columns else None
    )
    ip_col = 'ip_at_submit' if 'ip_at_submit' in df.columns else (
        'remote_ip' if 'remote_ip' in df.columns else None
    )
    country_col = 'country' if 'country' in df.columns else None
    region_col = 'region' if 'region' in df.columns else None
    city_col = 'city' if 'city' in df.columns else None
    course_col = 'course_name' if 'course_name' in df.columns else None

    def _format_time(val):
        if pd.isna(val):
            return ""
        try:
            if hasattr(val, 'to_pydatetime'):
                val = val.to_pydatetime()
            return val.strftime("%m/%d/%Y at %I:%M %p")
        except Exception:
            try:
                return pd.to_datetime(val).strftime("%m/%d/%Y at %I:%M %p")
            except Exception:
                return str(val)

    def _safe_str(x):
        return "" if pd.isna(x) else str(x)

    # Sort by submitted time if present
    if time_col and time_col in df.columns:
        df_sorted = df.sort_values(by=time_col)
    else:
        df_sorted = df.copy()

    # Build the document
    doc = Document()
    # Bold title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(f"{username} logs summary report")
    r_title.bold = True

    if course_col:
        for course in sorted(df_sorted[course_col].fillna('(Unknown Course)').unique(), key=lambda x: str(x).lower()):
            c_df = df_sorted[df_sorted[course_col].fillna('(Unknown Course)') == course]
            doc.add_paragraph("")
            # Bold 'Course Name:' label and course value
            p_course = doc.add_paragraph()
            r_label = p_course.add_run("Course Name: ")
            r_label.bold = True
            r_value = p_course.add_run(str(course))
            r_value.bold = True
            doc.add_paragraph("Assignments Submitted:")
            for _, row in c_df.iterrows():
                assignment_name = _safe_str(row[assign_col]) if assign_col else "(Unknown Assignment)"
                ip_val = _safe_str(row[ip_col]) if ip_col else ""
                submitted_val = _format_time(row[time_col]) if time_col else ""

                parts = []
                if city_col and not pd.isna(row.get(city_col)):
                    parts.append(str(row[city_col]))
                if region_col and not pd.isna(row.get(region_col)):
                    parts.append(str(row[region_col]))
                if country_col and not pd.isna(row.get(country_col)):
                    parts.append(str(row[country_col]))
                location_str = ", ".join([p for p in parts if p])

                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{assignment_name}")
                p = doc.add_paragraph(style='List Bullet 2')
                p.add_run(f"Submitted: {submitted_val}")
                p = doc.add_paragraph(style='List Bullet 2')
                p.add_run(f"IP Address: {ip_val}")
                p = doc.add_paragraph(style='List Bullet 2')
                p.add_run(f"IP Address Location: {location_str}")
    else:
        doc.add_paragraph("Assignments Submitted:")
        for _, row in df_sorted.iterrows():
            assignment_name = _safe_str(row[assign_col]) if assign_col else "(Unknown Assignment)"
            ip_val = _safe_str(row[ip_col]) if ip_col else ""
            submitted_val = _format_time(row[time_col]) if time_col else ""

            parts = []
            if city_col and not pd.isna(row.get(city_col)):
                parts.append(str(row[city_col]))
            if region_col and not pd.isna(row.get(region_col)):
                parts.append(str(row[region_col]))
            if country_col and not pd.isna(row.get(country_col)):
                parts.append(str(row[country_col]))
            location_str = ", ".join([p for p in parts if p])

            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{assignment_name}")
            p = doc.add_paragraph(style='List Bullet 2')
            p.add_run(f"Submitted: {submitted_val}")
            p = doc.add_paragraph(style='List Bullet 2')
            p.add_run(f"IP Address: {ip_val}")
            p = doc.add_paragraph(style='List Bullet 2')
            p.add_run(f"IP Address Location: {location_str}")

    doc.save(docx_path)
    print(f"  → Wrote submissions summary Word doc to {docx_path}")
def export_query(query_type: str, username: str, start_ts: datetime, end_ts: datetime, output_path: str) -> int:
    """
    Execute a parameterized SQL query and export results to Excel.
    
    Args:
        query_type: Type of query ('activity' or 'submissions')
        username: Canvas unique_id to filter
        start_ts: Start timestamp (EST/EDT, inclusive)
        end_ts: End timestamp (EST/EDT, exclusive)
        output_path: Path to output Excel file
    
    Returns:
        Number of rows exported
    """
    if query_type not in QUERY_TYPES:
        raise ValueError(f"Unknown query type: {query_type}. Must be one of: {', '.join(QUERY_TYPES.keys())}")
    
    if end_ts <= start_ts:
        raise ValueError("end timestamp must be greater than start timestamp")
    
    sql_path = QUERY_TYPES[query_type]
    sql = read_sql(sql_path)

    # Convert EST input to UTC for database query
    start_utc = start_ts.astimezone(ZoneInfo("UTC"))
    end_utc = end_ts.astimezone(ZoneInfo("UTC"))

    # Fetch data from database
    with connect_postgres() as conn:
        with conn.cursor() as cur:
            cur.execute(sql, {"username": username, "start_ts": start_utc, "end_ts": end_utc})
            rows = cur.fetchall()
            cols = [desc[0] for desc in cur.description]
    
    # Create DataFrame from results
    df = pd.DataFrame(rows, columns=cols)
    
    # Convert all datetime columns from UTC to EST/EDT
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            # Handle both timezone-aware and naive datetimes
            if df[col].dt.tz is None:
                df[col] = df[col].dt.tz_localize('UTC').dt.tz_convert('America/New_York').dt.tz_localize(None)
            else:
                df[col] = df[col].dt.tz_convert('America/New_York').dt.tz_localize(None)
    
    # Add IP geolocation if IP column exists
    ip_cols = [col for col in df.columns if col.lower() in ('ip', 'remote_ip', 'ip_at_submit')]
    
    if ip_cols:
        ip_col = ip_cols[0]
        
        # Get unique IPs and lookup locations
        unique_ips = df[ip_col].dropna().unique()
        print(f"  → Looking up locations for {len(unique_ips)} unique IPs...")
        
        # Batch lookup and create mapping
        location_map = {}
        for i, ip in enumerate(unique_ips, 1):
            if i % 10 == 0:  # Progress indicator every 10 IPs
                print(f"    {i}/{len(unique_ips)} IPs processed...")
            location_map[str(ip)] = get_ip_location(str(ip))
        
        # Map locations to all rows using vectorized operations
        df['country'] = df[ip_col].apply(lambda x: location_map.get(str(x), {}).get('country') if pd.notna(x) else None)
        df['region'] = df[ip_col].apply(lambda x: location_map.get(str(x), {}).get('region') if pd.notna(x) else None)
        df['city'] = df[ip_col].apply(lambda x: location_map.get(str(x), {}).get('city') if pd.notna(x) else None)
        df['isp'] = df[ip_col].apply(lambda x: location_map.get(str(x), {}).get('isp') if pd.notna(x) else None)
        
        # Reorder columns to place geolocation after IP
        ip_idx = df.columns.get_loc(ip_col)
        cols = [c for c in df.columns if c not in ['country', 'region', 'city', 'isp']]
        cols.insert(ip_idx + 1, 'country')
        cols.insert(ip_idx + 2, 'region')
        cols.insert(ip_idx + 3, 'city')
        cols.insert(ip_idx + 4, 'isp')
        df = df[cols]
    
    # Write to Excel using pandas (faster than openpyxl row-by-row)
    df.to_excel(output_path, index=False, sheet_name=query_type.capitalize())

    # If this is the submissions export, also write a Word summary
    if query_type == "submissions":
        write_submissions_summary_docx(df, username, start_ts, end_ts, output_path)

    return len(df)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Export Canvas activity logs and submissions from PostgreSQL")
    parser.add_argument("--username", required=True, help="Canvas login/unique_id to filter")
    parser.add_argument("--start", required=True, type=parse_dt, help="Start timestamp (EST)")
    parser.add_argument("--end", required=True, type=parse_dt, help="End timestamp (EST, exclusive)")
    parser.add_argument("--output-dir", required=False, help="Output directory (defaults to the username)")
    args = parser.parse_args(argv)

    try:
        # Resolve output directory (default to username if not provided)
        output_dir = args.output_dir or args.username
        os.makedirs(output_dir, exist_ok=True)
        
        # Get base name from output directory for file naming
        dir_basename = os.path.basename(os.path.normpath(output_dir))
        
        # Run both queries
        results = {}
        for query_type in ["activity", "submissions"]:
            output_file = os.path.join(output_dir, f"{dir_basename}-{query_type}.xlsx")
            print(f"Running {query_type} query...")
            count = export_query(query_type, args.username, args.start, args.end, output_file)
            results[query_type] = (count, output_file)
            print(f"  → Wrote {count} rows to {output_file}")
        
        print(f"\nCompleted successfully! Results in: {output_dir}/")
        return 0
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
